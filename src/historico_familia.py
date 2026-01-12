from docx import Document
from pathlib import Path
from database import buscar_visitas_familia 

BASE_DIR = Path(__file__).resolve().parent.parent
REPORTS_DIR = BASE_DIR / "reports"
REPORTS_DIR.mkdir(exist_ok=True)


def gerar_relatorio_historico(familia):
    visitas = buscar_visitas_familia(familia)

    if not visitas:
        print(f"Nenhuma visita encontrada para {familia}")
        return

    doc = Document()

    doc.add_heading(f"RELATÓRIO HISTÓRICO – {familia}", level=1)
    doc.add_paragraph("")

    tecnico_referencia = visitas[-1][1]
    doc.add_paragraph(f"Técnico de referência: {tecnico_referencia}")
    doc.add_paragraph("")

    for data_visita, tecnico, descricao, risco, encaminhamentos in visitas:
        data_br = data_visita.split("-")[2] + "/" + data_visita.split("-")[1] + "/" + data_visita.split("-")[0]

        doc.add_heading(f"VISITA – {data_br}", level=2)
        doc.add_paragraph(f"Risco: {risco}")
        doc.add_paragraph("")
        doc.add_paragraph("Descrição da visita:")
        doc.add_paragraph(descricao)
        doc.add_paragraph("")
        doc.add_paragraph("Encaminhamentos:")
        doc.add_paragraph(encaminhamentos if encaminhamentos else "Não informado")
        doc.add_paragraph("\n" + "-" * 40 + "\n")

    nome_arquivo = f"Historico_{familia.replace(' ', '_')}.docx"
    doc.save(REPORTS_DIR / nome_arquivo)

    print(f"Relatório histórico gerado com sucesso: {nome_arquivo}")


if __name__ == "__main__":
    familia = input("Digite o nome da família: ").strip()
    gerar_relatorio_historico(familia)
